#!/usr/bin/env python3
"""Build the ACADIA 2026 paper from the template DOCX.

Run:
    python3.10 cases/apartment_daytime_cooled/build_paper.py
"""
from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CASE = Path(__file__).resolve().parent
TEMPLATE = CASE / "ACADIA_2026_Publication_Templates+Guidelines" / "ACADIA_2026_FullPapers_StyleGuide.docx"
OUTPUT = CASE / "paper_draft.docx"
FIGS = CASE / "out" / "figures"

# ── Image paths ────────────────────────────────────────────────────

IMG = {
    "pipeline":     FIGS / "fig0_framework_pipeline_hires.png",
    "occupant_map": CASE / "Physics_Occupant.png",
    "json_snippet": FIGS / "fig_json_snippet_hires.png",
    "plan":         FIGS / "fig1_apartment_plan_hires.png",
    "matrix":       FIGS / "fig2_design_matrix_hires.png",
    "span":         FIGS / "fig3_comfort_span_hires.png",
    "decomp":       FIGS / "fig4_heat_loss_decomposition_hires.png",
    "scatter":      FIGS / "fig5_conv_rad_scatter_hires.png",
    "office_story": CASE.parent / "office_topology_compare" / "out" / "figures" / "fig_office_topology_story_with_box_sidecar_hires.png",
}

# ═══════════════════════════════════════════════════════════════════
# Paper content
# ═══════════════════════════════════════════════════════════════════

TITLE = (
    "Hidden Spatial Intelligence: Revealing Thermal Comfort "
    "Heterogeneity Through Occupant-Level Heat-Balance Assessment"
)

ABSTRACT = (
    "Floor plans carry thermal consequences that conventional representations cannot show. "
    "A south-facing bedroom occupied by an older resident can produce a fundamentally different "
    "thermal experience than a north-facing studio used by a younger desk worker, even under "
    "the same setpoint. This paper presents a spatial comfort assessment framework that makes "
    "that hidden differentiation legible by distributing individually characterized occupant "
    "agents across a gridded floor plan and computing ISO 7730 heat-balance decomposition per "
    "agent, turning anticipated design conditions into a population-level thermal field rather "
    "than a zone-level average. Applied to a 24-room apartment under 120 scenarios (five air "
    "temperatures, four humidity levels, six occupant cohorts spanning demographic and "
    "behavioral variation), the framework produces 2,520 agent-level records and per-agent "
    "JSON snapshots carrying full convective, radiative, evaporative, and respiratory heat-loss "
    "breakdowns. In the apartment case, room-level PMV spans range from roughly one to nearly "
    "four units within a single plan under uniform setpoints. A secondary office counterfactual "
    "on the same footprint shows that partition topology does not neutralize discomfort so much "
    "as redistribute it: at 22\u00b0C / 80% RH, the open-office interpretation widens work-cell "
    "PMV spread for every cohort while leaving mean sensation nearly unchanged and, for most "
    "cohorts, slightly colder than the cellular layout. Orientation-driven mean radiant "
    "temperature offsets shift "
    "radiative heat loss by about 5 W/m\u00b2, while cohort configurations involving clothing and "
    "activity shift convective and radiative loss by similar magnitudes. Spatial organization "
    "is not a passive container for thermal conditions. It is an active participant in who gets "
    "comfortable and who does not."
)

KEYWORDS = "thermal comfort, spatial intelligence, heat-balance decomposition, PMV, floor plan assessment"

SECTIONS = [
    # (heading_level, heading_number, heading_text, paragraphs)
    # heading_level: 1 = Header A, 2 = Header B, 0 = unnumbered Header A
    # paragraphs: list of strings (each is one paragraph)

    (0, "", "Abstract", [ABSTRACT]),

    (1, "1", "Introduction", [

        "Architectural design is always directed toward occupants, yet early environmental "
        "reasoning often treats those occupants as generic. Standard workflows reduce a "
        "building to conditioned zones and report average temperatures and aggregate loads. "
        "The embedded assumption is that if the thermostat is set appropriately, the zone is "
        "comfortable and the people inside it are effectively interchangeable.",

        "That assumption does not survive contact with an actual floor plan. Under the same "
        "25\u00b0C setpoint, mean radiant temperature varies by orientation, air velocity varies "
        "by room function, and occupants themselves modulate the thermal balance through their "
        "age, body mass, clothing, and activity. A south-facing bedroom occupied by an older "
        "resident in light clothing produces a fundamentally different thermal experience than "
        "a north-facing kitchen with a younger person standing and cooking. The plan encodes "
        "this differentiation. Conventional tools and drawings do not surface it.",

        "The gap this creates is not merely technical. It is representational. Architects make "
        "decisions about room assignment, partition placement, glazing, and orientation during "
        "early design, long before any occupant sets foot in the building. Those decisions carry "
        "thermal consequences for specific kinds of people, yet the standard zone-level model "
        "collapses those consequences into a single average. The plan communicates spatial "
        "intent. It does not communicate who will be comfortable and who will not.",

        "This paper presents a framework that closes that representational gap. By distributing "
        "individually characterized occupant agents across a gridded floor plan and computing "
        "ISO 7730 heat-balance decomposition per agent, the framework translates anticipated "
        "design conditions into a population-level thermal field readable at the resolution "
        "where design decisions happen. The result is not an optimization dashboard. It is a "
        "design-facing analytical drawing that makes the thermal consequences of spatial "
        "decisions visible before construction.",

        "The framework combines three input layers\u2014spatial geometry, room-level "
        "environmental conditions, and individually characterized occupant agents\u2014and "
        "produces per-agent JSON snapshots carrying full heat-loss decomposition alongside "
        "demographic, spatial, and preference context. Applied to a 24-room apartment under "
        "120 scenarios and extended to an office topology counterfactual on the same footprint, "
        "the case study demonstrates that spatial organization and body configuration produce "
        "comparable magnitudes of thermal heterogeneity, and that partition topology widens and "
        "clarifies perimeter-to-core inequality more than it resolves the discomfort itself.",

        "The contribution is not a new dynamic simulation method. It is a lightweight, "
        "deterministic representational instrument for early design: one that reveals where a "
        "given plan becomes population-conditioned and spatially fragile, where comfort outcomes "
        "diverge sharply across occupant profiles and room positions, and where room assignment, "
        "orientation, partition placement, glazing specification, and environmental setpoints "
        "carry consequences that conventional zone-level analysis does not show.",
    ]),

    (1, "2", "Related Work", [

        "Fanger\u2019s Predicted Mean Vote model (Fanger 1970), codified in ISO 7730 (ISO 2005) "
        "and ASHRAE Standard 55 (ASHRAE 2020), remains the primary analytical framework for "
        "assessing steady-state thermal comfort in mechanically conditioned buildings. The model "
        "resolves the six-variable heat balance\u2014air temperature, mean radiant temperature, "
        "relative humidity, air velocity, metabolic rate, and clothing insulation\u2014into a "
        "single-index prediction of occupant thermal sensation. Its strength lies in its "
        "traceability: every intermediate quantity, from clothing surface temperature to "
        "convective film coefficient, is computable and auditable.",

        "In practice, however, the PMV model is typically applied at the zone level. The CBE "
        "Thermal Comfort Tool (Tartarini et al. 2020) provides an accessible web-based calculator "
        "for evaluating individual conditions but does not map comfort across spatial layouts. "
        "Building performance simulation platforms such as EnergyPlus (Crawley et al. 2008) produce "
        "zone-level temperature and humidity outputs that feed PMV calculations, but the spatial "
        "resolution stops at the zone boundary\u2014room-to-room gradients within a single dwelling "
        "are collapsed. Higher-fidelity approaches including computational fluid dynamics and "
        "zonal models (Megri and Haghighat 2007) can resolve spatial thermal fields but require "
        "substantial modeling effort and computational time, placing them outside the workflow "
        "of most architectural designers.",

        "Agent-based modeling has a longer history in architectural research, primarily for "
        "pedestrian movement simulation, evacuation analysis, and space utilization studies "
        "(Yan and Kalay 2006). Agents serve as spatial probes: by distributing synthetic "
        "occupants across a design and tracking their experience, patterns emerge that static "
        "analysis cannot reveal. Extending the agent metaphor to thermal comfort\u2014placing "
        "thermally characterized agents on a floor plan and computing their individual heat "
        "balance\u2014is the approach this paper takes.",

        "Adaptive comfort models (de Dear and Brager 1998; Nicol and Humphreys 2002) have "
        "broadened the theoretical frame beyond PMV by incorporating occupant behavioral "
        "adaptation, but these models are most applicable to naturally ventilated buildings "
        "and do not decompose the heat balance into spatial components.",

        "More broadly, computational architecture has developed a tradition of making "
        "latent spatial performance visible through analysis feedback. Space Syntax "
        "(Hillier and Hanson 1984) revealed measurable connectivity and integration patterns "
        "in plans that were invisible to inspection alone. Daylighting analysis, pedestrian "
        "simulation, and acoustic mapping followed the same pattern: translating a spatial "
        "design into a performance field that designers can read and act on. The framework "
        "presented here extends this tradition to thermal comfort at the individual occupant "
        "level. The plan is treated not as a container to be conditioned uniformly but as a "
        "differentiated field whose spatial organization produces heterogeneous thermal "
        "outcomes that are recoverable through computation.",
    ]),

    (1, "3", "Methods", []),

    (2, "3.1", "Framework Overview", [
        "The framework operates on three input layers (Figure 1). The first is spatial "
        "geometry: a grid-based floor plan with rooms defined as contiguous sets of cells, "
        "bounded by walls and connected by doors and openings. The second is environmental "
        "conditions: each room carries a static set of anticipated design parameters "
        "(air temperature, mean radiant temperature, relative humidity, air velocity, "
        "illuminance, and background noise level) assigned from design specifications rather "
        "than simulated. The third is occupant agents: individually placed persons "
        "characterized by demographics, clothing insulation, activity level, and preferences.",

        "The framework is structured to accept any floor plan that can be segmented into "
        "rooms with assigned environmental conditions. Geometry enters through an assisted "
        "PDF import tool that traces architectural drawings onto the working grid, or through "
        "a direct Rhino export pathway. The apartment presented in this paper is one instance; "
        "the pipeline is structured to accommodate other typologies (offices, healthcare "
        "facilities, educational buildings) provided room-level conditions can be reasonably "
        "anticipated, though these have not been validated in the current study. An online "
        "version of the geometry import and assessment tools is available (URL withheld for "
        "anonymous review).",

        "These three layers feed a deterministic heat-balance engine that computes "
        "ISO 7730 PMV and full thermal decomposition per agent. The computation is not a "
        "simulation in the dynamic sense: room states do not evolve over time, and occupant "
        "heat gains do not feed back into room conditions. The tool evaluates the instantaneous "
        "thermal balance at each agent\u2019s position given the specified conditions. This "
        "is a calculation, not a simulation. The distinction is deliberate: the framework runs "
        "instantly and produces fully traceable results, enabling rapid iteration during design.",

        "The per-agent output is a structured JSON snapshot (Figure 2) carrying the full "
        "context of one occupant\u2019s experience: environment, spatial position, computed "
        "thermal state with heat-loss decomposition, demographics, and preferences. The "
        "record is computationally grounded in empirical heat-transfer relationships while "
        "remaining structured for downstream consumption by analytical or interpretive tools.",
    ]),

    (2, "3.2", "Thermal Computation", [
        "The thermal core implements the standard ISO 7730 iterative procedure. Metabolic "
        "rate is estimated from occupant demographics using the Mifflin-St Jeor basal "
        "metabolic rate equation (Mifflin et al. 1990) scaled by a Du Bois body surface "
        "area estimate (Du Bois and Du Bois 1916), then adjusted by an activity multiplier "
        "corresponding to seated, standing, or walking postures. Clothing insulation converts "
        "to the intrinsic thermal resistance of the clothing ensemble.",

        "The iterative solver finds the clothing surface temperature at which the heat-balance "
        "equation is satisfied, then computes each outgoing heat-transfer pathway independently: "
        "convective loss from the clothing surface to ambient air, radiative exchange between "
        "the clothing surface and surrounding surfaces at the mean radiant temperature, "
        "evaporative loss through skin diffusion and regulatory sweating, and respiratory "
        "heat loss through convection and evaporation in the lungs. These four components\u2014"
        "reported in watts per square meter of body surface area\u2014are preserved in the "
        "agent snapshot alongside the aggregate PMV and PPD values. It is this decomposition "
        "that enables the spatial analysis presented in the results: by examining which heat-loss "
        "channel varies across rooms or across occupant profiles, the framework reveals the "
        "physical mechanism driving comfort heterogeneity, not merely its summary index.",
    ]),

    (2, "3.3", "Case Study Design", [
        "To demonstrate the framework on a realistic plan, the case study uses a 24-room "
        "apartment traced from architectural drawings through the PDF import pipeline. The "
        "plan is discretized onto a 32\u00d724 meter grid at one-meter cell resolution "
        "(Figure 3). Twenty-one rooms are occupied; three circulation zones are excluded. "
        "The apartment was selected as a reasonably generic residential typology; the "
        "framework is equally applicable to any plan that can be segmented into rooms with "
        "assigned conditions.",

        "Environmental conditions are assigned deterministically rather than simulated. Air "
        "temperature and relative humidity are set uniformly across the apartment for each "
        "scenario, representing a centrally conditioned dwelling. Mean radiant temperature is "
        "differentiated by solar exposure: south-facing rooms receive a +1.5 K offset above "
        "air temperature, east- and west-facing rooms receive +0.75 K, and north-facing or "
        "interior rooms receive no offset. Air velocity is assigned by room function: 0.08 m/s "
        "for bedrooms, 0.12 m/s for day spaces, 0.15 m/s for the kitchen, 0.08 m/s for bath "
        "and laundry spaces, and 0.10 m/s for the remaining support rooms. These values "
        "represent typical conditions in a mechanically cooled apartment during daytime "
        "operation.",

        "Six occupant cohorts are defined to span a range of demographic and behavioral "
        "profiles: a mixed-gender young cohort (ages 25\u201345), a male-only control, an "
        "age-shifted cohort (+50 years), a higher-clothing sedentary cohort, a lighter-clothing "
        "mobile cohort, and a higher-BMI cohort. Within each cohort, individual agent demographics "
        "are generated from a seeded random process producing variation in age, height, weight, "
        "and sex assignment across the 21 occupied rooms. Clothing insulation and activity level "
        "are assigned by room function and modified by cohort definition.",

        "The climate sweep covers five air temperatures (20.0, 22.0, 23.5, 25.0, 28.0\u00b0C) "
        "and four relative humidity levels (35, 55, 75, 80%), producing 20 environmental "
        "conditions per cohort and 120 scenarios total. Each scenario places 21 agents\u2014one "
        "per occupied room\u2014yielding 2,520 individual agent records with full heat-balance "
        "decomposition.",
    ]),

    (1, "4", "Results", []),

    (2, "4.1", "Spatial Comfort Heterogeneity Under Varying Design Conditions", [
        "Figure 4 presents the central result as a three-by-three matrix of apartment plans "
        "colored by occupant PMV. Columns correspond to three occupant cohorts (higher clothing, "
        "young mixed, and age-shifted), and rows correspond to three air temperature setpoints "
        "(20, 25, and 28\u00b0C), all at 55% relative humidity. Reading horizontally across "
        "any row reveals the body-configuration effect: at the same temperature, the higher-"
        "clothing cohort shows consistently milder thermal stress than the age-shifted cohort, "
        "with mean PMV differences of roughly 1.1 to 2.4 units across the shown setpoints. "
        "Reading vertically down any column "
        "reveals the design-condition effect: raising the setpoint from 20 to 28\u00b0C shifts "
        "the entire comfort landscape from deep cold stress to near-neutral or mildly warm "
        "conditions for some cohorts while older occupants remain cooler overall.",

        "Critically, the spatial pattern within the plan is visible at every combination. "
        "South-facing rooms consistently appear warmer than north-facing rooms due to the "
        "mean radiant temperature offset, producing a legible thermal gradient across the "
        "plan that persists regardless of cohort or setpoint. The framework makes this "
        "gradient\u2014encoded in the spatial organization but invisible in conventional "
        "plan representations\u2014directly visible.",
    ]),

    (2, "4.2", "Quantified Comfort Span", [
        "Figure 5 presents the comfort span as a strip chart plotting each occupant\u2019s PMV "
        "across five air temperatures for three cohorts. Each dot represents one agent in one "
        "room; marker shape encodes solar exposure. Several patterns are immediately apparent.",

        "First, within any cohort column at any temperature, the vertical spread of dots spans "
        "roughly one to nearly four PMV units depending on cohort and setpoint. This is the "
        "spatial component of the heterogeneity: "
        "the floor plan produces meaningfully different thermal conditions room to room even "
        "under a uniform setpoint. Second, between cohort columns at the same temperature, "
        "the cohort means separate by roughly one to more than two PMV units. The spatial and "
        "demographic "
        "effects compound. Third, south-facing rooms (triangular markers) consistently plot "
        "above north-facing rooms (circular markers) within each cohort, confirming that the "
        "MRT gradient from orientation is a persistent and legible spatial signal. Fourth, the "
        "span narrows as temperature increases toward the comfort zone, but it does not disappear, "
        "indicating that the spatial intelligence encoded in "
        "the plan is a structural feature of the design, not an artifact of extreme conditions.",
    ]),

    (2, "4.3", "Heat-Balance Decomposition", [
        "Figure 6 presents the heat-loss decomposition for the young mixed cohort at 25\u00b0C "
        "and 55% relative humidity. Each horizontal bar represents one occupied room, sorted "
        "by total heat loss and divided into convective, radiative, evaporative, and respiratory "
        "components. Exposure zone is indicated by color chips on the left margin; mean radiant "
        "temperature and air velocity are annotated on the right.",

        "Total heat loss ranges from 58 to 69 W/m\u00b2 across the 21 rooms. The variation is "
        "driven primarily by two spatial factors. The radiative component drops approximately "
        "five watts per square meter from north-facing rooms (mean radiant temperature equal to "
        "air temperature) to south-facing rooms (mean radiant temperature 1.5 K above air "
        "temperature). This is the thermal signature of solar exposure encoded in the plan\u2019s "
        "orientation. The convective component peaks in the kitchen at 28.1 W/m\u00b2, driven by "
        "the higher assigned air velocity of 0.15 m/s, compared to 22\u201323 W/m\u00b2 in bedrooms "
        "at 0.08 m/s. Evaporative and respiratory losses are comparatively uniform across rooms, "
        "contributing approximately 11 and 4 W/m\u00b2 respectively, and varying primarily with "
        "individual metabolic rate rather than room conditions.",
    ]),

    (1, "5", "Discussion", []),

    (2, "5.1", "Comparable Magnitudes of Spatial and Body Effects", [
        "The heat-balance decomposition makes it possible to ask a design question that "
        "conventional comfort analysis usually hides: does the spatial organization of a plan "
        "matter as much as the people inside it? Figure 7 plots convective against radiative "
        "heat loss for each room in three cohorts at 25\u00b0C. Color encodes cohort; marker "
        "shape encodes solar exposure. A diagonal reference marks equal convective and "
        "radiative loss.",

        "Three findings emerge. First, cohort clusters separate clearly along both axes. "
        "At 25\u00b0C / 55% RH, the higher-clothing sedentary cohort shifts mean convective and "
        "radiative loss down by about 4.8 and 5.0 W/m\u00b2 respectively relative to the young "
        "mixed baseline. The body acts as a thermal shell modulating multiple outgoing "
        "pathways simultaneously.",

        "Second, within each cohort cluster, south-facing rooms (triangles) pull down by "
        "approximately 5 W/m\u00b2 in radiative loss relative to north-facing rooms (circles), "
        "while convective loss shifts only slightly. The orientation-driven MRT gradient "
        "modulates the radiative pathway directly without proportionally affecting convection.",

        "The central finding is that these two effects are comparable in magnitude. A 1.5 K "
        "MRT offset from room orientation and a cohort configuration change centered on higher "
        "clothing and reduced activity each displace the heat-transfer balance by roughly "
        "5 W/m\u00b2. Design decisions about "
        "spatial organization carry thermal consequences of the same order as the occupant "
        "variables designers typically consider uncontrollable. The plan is not thermally "
        "neutral. It is an active participant in the comfort equation.",
    ]),

    (2, "5.2", "Limitations", [
        "Several hard limitations bound the current results and should be stated plainly. "
        "All room environmental conditions are static and hand-authored: air temperature, "
        "mean radiant temperature, humidity, and air velocity are assigned per room from "
        "anticipated design values, not derived from simulation. Occupant heat gains do not "
        "feed back into room state. A room occupied by twenty people will report the same "
        "air temperature as an empty one. This is the most consequential simplification.",

        "Grid resolution is one meter. This is adequate for room-level comfort mapping but "
        "does not resolve microclimate gradients near windows, radiant panels, or diffusers. "
        "Mean radiant temperature is differentiated by exposure classification (south, "
        "east/west, north/interior) rather than by geometric view-factor calculation, and "
        "no solar penetration or radiant asymmetry is modeled. The MRT offsets used in this "
        "study (+1.5 K south, +0.75 K east/west) are reasonable assumptions for a daytime-"
        "cooled apartment but are not derived from radiation modeling.",

        "The PMV model itself is validated for steady-state conditions within a bounded "
        "input range. Values approaching \u00b13 should be interpreted as indicators of strong "
        "thermal stress rather than precise predictions of occupant sensation. Some plotted PMV "
        "figures in this paper are clamped on the cold side for legibility, but the stored batch "
        "outputs retain the full computed range.",
    ]),

    (2, "5.3", "Design Interventions and Future Work", [
        "Within these limitations, the framework supports the kind of iterative judgment that "
        "characterizes early design. Modifying wall partitions changes room geometry and zone "
        "assignment. Increasing glazing area or changing facade type alters the mean radiant "
        "temperature assigned to affected rooms. Adjusting mechanical ventilation changes "
        "per-room air velocity. Each intervention propagates instantly through the heat-balance "
        "engine, enabling direct comparison of comfort outcomes across plan variants.",

        "The office counterfactual suggests that topology can redistribute and sometimes deepen "
        "cold-side discomfort even when it does not materially improve mean comfort (Figure 8). "
        "When the same envelope is reinterpreted as either a cellular office or an open-plan "
        "office under a cool-humid 22\u00b0C / 80% RH condition, the cohort means shift only "
        "marginally, with three of the four compared cohorts moving slightly colder, while the "
        "work-cell field stretches substantially. For a realistic mixed office cohort, "
        "interpolated work-cell PMV range increases from 0.216 in the cellular layout to 0.662 "
        "in the open layout; for a default 35-year-old male baseline, from 0.179 to 0.529; for "
        "a realistic male cohort, from 0.182 to 0.537; and for a lighter-clothed female cohort, "
        "from 0.245 to 0.772. The design intelligence here is not that open plans solve the "
        "comfort problem. It is that topology determines whether perimeter-to-core inequality is "
        "averaged away into rooms or exposed as a continuous field that can be read, assigned, "
        "and redesigned.",

        "The occupant model accepts arbitrary agent definitions beyond the six cohorts "
        "presented here, including high-density layouts and sparse occupancy patterns. "
        "Scheduled events allow agents to move between rooms and change clothing or activity "
        "over time. Future work will address transient room-level coupling, view-factor-based "
        "MRT calculation, and probabilistic agent placement to quantify comfort robustness "
        "across occupancy scenarios. The data model already carries fields for illuminance, "
        "background noise, and social density alongside the thermal variables exercised here, "
        "positioning the framework for multi-channel environmental quality assessment.",
    ]),

    (1, "6", "Conclusion", [
        "Much thermal analysis still implies a generic occupant and a generic zone. This paper "
        "argues that those abstractions hide the consequences of spatial decisions for different "
        "bodies in different rooms. By distributing individually characterized agents across a "
        "floor plan, the framework makes hidden thermal intelligence legible before "
        "construction.",

        "In the 24-room apartment case, that hidden intelligence appears as persistent comfort "
        "heterogeneity ranging from roughly one to nearly four PMV units within a single plan "
        "under uniform setpoints. The heat-balance decomposition shows why. Spatial factors\u2014"
        "primarily orientation-driven mean radiant temperature gradients\u2014and body-"
        "configuration factors\u2014most visibly clothing and activity configuration in this "
        "study\u2014shift heat loss by comparable magnitudes on the order of five watts per "
        "square meter. The architectural plan is therefore not a passive container for thermal "
        "conditions. Its spatial decisions carry thermal consequences on the same order as "
        "occupant variables designers do not typically control.",

        "The office counterfactual sharpens this claim. Reusing the same footprint as a cellular "
        "or open office does not improve comfort simply by removing partitions; at 22\u00b0C, many "
        "occupants remain cold in both layouts. Open plan slightly worsens mean sensation for "
        "most cohorts and widens work-cell spread for all of them. What changes most is the "
        "distribution and legibility of discomfort. Open topology stretches the perimeter-to-core "
        "gradient into a continuous field, while cellular topology collapses that same gradient "
        "back into room averages.",

        "What the framework ultimately offers is not another optimization target but a new kind "
        "of design feedback: the ability to see, before a building is built, which occupants a "
        "given spatial configuration advantages, which it penalizes, and why. The holistic data "
        "model already carries lighting, acoustic, and preference fields alongside the thermal "
        "channel exercised here, positioning the framework for extension to multi-channel "
        "environmental quality assessment at the same spatial resolution.",
    ]),
]

REFERENCES = [
    'ASHRAE. 2020. ANSI/ASHRAE Standard 55-2020: Thermal Environmental Conditions for Human Occupancy. Atlanta: ASHRAE.',
    'Crawley, Drury B., Jon W. Hand, Micha\u00ebl Kummert, and Brent T. Griffith. 2008. \u201cContrasting the Capabilities of Building Energy Performance Simulation Programs.\u201d Building and Environment 43 (4): 661\u2013673.',
    'de Dear, Richard J. and Gail S. Brager. 1998. \u201cDeveloping an Adaptive Model of Thermal Comfort and Preference.\u201d ASHRAE Transactions 104 (1): 145\u2013167.',
    'Du Bois, Delafield and Eugene F. Du Bois. 1916. \u201cA Formula to Estimate the Approximate Surface Area if Height and Weight Be Known.\u201d Archives of Internal Medicine 17 (6): 863\u2013871.',
    'Fanger, Poul O. 1970. Thermal Comfort: Analysis and Applications in Environmental Engineering. Copenhagen: Danish Technical Press.',
    'Fountain, Marc and Charles Huizenga. 1995. \u201cA Thermal Sensation Model for Use by the Engineering Profession.\u201d ASHRAE Technical Data Bulletin 11 (1): 1\u201312.',
    'Hillier, Bill and Julienne Hanson. 1984. The Social Logic of Space. Cambridge: Cambridge University Press.',
    'Holopainen, Riikka. 2012. \u201cA Human Thermal Model for Improved Thermal Comfort.\u201d VTT Science 23. Espoo: VTT Technical Research Centre of Finland.',
    'ISO. 2005. ISO 7730:2005 Ergonomics of the Thermal Environment\u2014Analytical Determination and Interpretation of Thermal Comfort Using Calculation of the PMV and PPD Indices and Local Thermal Comfort Criteria. Geneva: International Organization for Standardization.',
    'Megri, Ahmed Cherif and Fariborz Haghighat. 2007. \u201cZonal Modeling for Simulating Indoor Environment of Buildings: Review, Recent Developments, and Applications.\u201d HVAC&R Research 13 (6): 887\u2013905.',
    'Mifflin, Mark D., Sachiko T. St Jeor, Lisa A. Hill, Barbara J. Scott, Sandra A. Daugherty, and Young O. Koh. 1990. \u201cA New Predictive Equation for Resting Energy Expenditure in Healthy Individuals.\u201d American Journal of Clinical Nutrition 51 (2): 241\u2013247.',
    'Nicol, J. Fergus and Michael A. Humphreys. 2002. \u201cAdaptive Thermal Comfort and Sustainable Thermal Standards for Buildings.\u201d Energy and Buildings 34 (6): 563\u2013572.',
    'Tartarini, Federico, Stefano Schiavon, Tyler Cheung, and Tyler Hoyt. 2020. \u201cCBE Thermal Comfort Tool: Online Tool for Thermal Comfort Calculations and Visualizations.\u201d SoftwareX 12: 100563.',
    'Yan, Wei and Yehuda E. Kalay. 2006. \u201cSimulating the Behavior of Users in Built Environments.\u201d Journal of Architectural and Planning Research 23 (4): 289\u2013310.',
]

FIGURE_CAPTIONS = {
    1: "Figure 1.  Workflow from plan authoring to batch outputs. PDF- or Rhino-derived geometry is simplified into a grid-based template, enriched with room conditions and occupant cohorts, then evaluated by thermBAL\u2019s deterministic heat-balance engine to produce PMV maps, comparative plots, and per-agent JSON snapshots.",
    2: "Figure 2.  Occupant perception map showing the structured snapshot for one agent. The record carries environment, spatial context, computed thermal state with heat-loss decomposition, agent demographics, and subjective preferences.",
    3: "Figure 3.  Apartment plan with 24 rooms classified by solar exposure. South-facing rooms receive a mean radiant temperature offset of +1.5 K; east/west rooms receive +0.75 K; north and interior rooms receive no offset.",
    4: "Figure 4.  Same spatial design under varying conditions and bodies. Columns: three occupant cohorts. Rows: three air temperature setpoints (20, 25, 28\u00b0C). Room color encodes occupant PMV clamped to [\u22123, +1.5]. All scenarios at 55% RH.",
    5: "Figure 5.  Comfort span across design conditions. Each dot is one occupant; horizontal bars show cohort mean. Marker shape encodes solar exposure. Depending on cohort and setpoint, the within-plan PMV spread ranges from roughly 1 to nearly 4 units and narrows toward the comfort zone without disappearing.",
    6: "Figure 6.  Heat-loss decomposition for the young mixed cohort at 25\u00b0C / 55% RH. Stacked bars show convective, radiative, evaporative, and respiratory components per room, sorted by total heat loss. Exposure chips on left; room conditions on right.",
    7: "Figure 7.  Convective versus radiative heat loss per occupant for three cohorts at 25\u00b0C. Marker shape encodes exposure. Spatial factors (orientation) and cohort configuration effects (especially clothing and activity) shift heat loss by comparable magnitudes on the order of 5 W/m\u00b2.",
    8: "Figure 8.  Office topology counterfactual on the shared footprint at 22\u00b0C / 80% RH. Left: the same footprint rendered as a cellular office and as an open office for three representative cohorts; color encodes interpolated PMV across occupied work areas, while support rooms remain muted. Right: paired boxplots show actual work-cell PMV distributions for four cohorts on a shared axis (graphite = cellular, moss = open). The open-plan version widens PMV spread for every cohort and typically nudges sensation colder, making perimeter-to-core thermal inequality more legible rather than eliminating it.",
}

ACKNOWLEDGEMENTS = (
    "Figure generation scripts and portions of the manuscript were developed with assistance "
    "from Claude Code (Anthropic, Claude Opus 4.6, 2026), a generative AI coding tool. The "
    "authors verified all computational outputs against the underlying ISO 7730 implementation "
    "and reviewed all AI-generated text for accuracy and completeness. The thermal comfort engine, "
    "floor-plan models, and batch study infrastructure were developed by the authors."
)


# ═══════════════════════════════════════════════════════════════════
# Document builder
# ═══════════════════════════════════════════════════════════════════

def build():
    doc = Document(str(TEMPLATE))

    # Clear template body (keep styles)
    for _ in range(len(doc.paragraphs)):
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    for t in doc.tables:
        t._element.getparent().remove(t._element)

    # ── Helpers ──
    def heading_a(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.bold = True; run.font.size = Pt(12); run.font.name = "Times New Roman"
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.space_after = Pt(6)

    def heading_b(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.font.size = Pt(12); run.font.name = "Times New Roman"
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after = Pt(6)

    def body(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12); run.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)

    def figure(img_key, caption, width=Inches(6.5)):
        path = IMG[img_key]
        if not path.exists():
            body(f"[IMAGE NOT FOUND: {path.name}]")
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=width)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        # Caption
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.font.size = Pt(10); run.font.name = "Times New Roman"
        run.italic = True
        cap.paragraph_format.space_after = Pt(12)

    # ═════════════════════════════════════════════════════════════
    # Title
    # ═════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True; run.font.size = Pt(14); run.font.name = "Times New Roman"
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(48)

    # ═════════════════════════════════════════════════════════════
    # Abstract
    # ═════════════════════════════════════════════════════════════
    heading_a("Abstract")
    body(ABSTRACT)
    p = doc.add_paragraph()
    run = p.add_run(f"Keywords: {KEYWORDS}")
    run.font.size = Pt(12); run.font.name = "Times New Roman"
    run.italic = True
    p.paragraph_format.space_after = Pt(12)

    # ═════════════════════════════════════════════════════════════
    # 1. Introduction
    # ═════════════════════════════════════════════════════════════
    heading_a("1 Introduction")
    for text in _section_paras("1"):
        body(text)

    # ═════════════════════════════════════════════════════════════
    # 2. Related Work
    # ═════════════════════════════════════════════════════════════
    heading_a("2 Related Work")
    for text in _section_paras("2"):
        body(text)

    # ═════════════════════════════════════════════════════════════
    # 3. Methods
    # ═════════════════════════════════════════════════════════════
    heading_a("3 Methods")

    heading_b("3.1 Framework Overview")
    for text in _section_paras("3.1"):
        body(text)

    # → Figure 1: Framework pipeline
    figure("pipeline",
           FIGURE_CAPTIONS[1],
           width=Inches(6.5))

    # → Figure 2: Occupant perception map (data model)
    figure("occupant_map",
           FIGURE_CAPTIONS[2],
           width=Inches(5.0))

    heading_b("3.2 Thermal Computation")
    for text in _section_paras("3.2"):
        body(text)

    heading_b("3.3 Case Study Design")
    for text in _section_paras("3.3"):
        body(text)

    # → Figure 3: Apartment plan
    figure("plan",
           FIGURE_CAPTIONS[3],
           width=Inches(6.5))

    # ═════════════════════════════════════════════════════════════
    # 4. Results
    # ═════════════════════════════════════════════════════════════
    heading_a("4 Results")

    heading_b("4.1 Spatial Comfort Heterogeneity Under Varying Design Conditions")
    for text in _section_paras("4.1"):
        body(text)

    # → Figure 4: Design matrix (3×3)
    figure("matrix",
           FIGURE_CAPTIONS[4],
           width=Inches(6.5))

    heading_b("4.2 Quantified Comfort Span")
    for text in _section_paras("4.2"):
        body(text)

    # → Figure 5: Comfort span strip chart
    figure("span",
           FIGURE_CAPTIONS[5],
           width=Inches(6.5))

    heading_b("4.3 Heat-Balance Decomposition")
    for text in _section_paras("4.3"):
        body(text)

    # → Figure 6: Heat loss stacked bars
    figure("decomp",
           FIGURE_CAPTIONS[6],
           width=Inches(6.0))

    # ═════════════════════════════════════════════════════════════
    # 5. Discussion
    # ═════════════════════════════════════════════════════════════
    heading_a("5 Discussion")

    heading_b("5.1 Comparable Magnitudes of Spatial and Body Effects")
    for text in _section_paras("5.1"):
        body(text)

    # → Figure 7: Conv-rad scatter
    figure("scatter",
           FIGURE_CAPTIONS[7],
           width=Inches(5.0))

    heading_b("5.2 Limitations")
    for text in _section_paras("5.2"):
        body(text)

    heading_b("5.3 Design Interventions and Future Work")
    for text in _section_paras("5.3"):
        body(text)

    figure("office_story",
           FIGURE_CAPTIONS[8],
           width=Inches(6.5))

    # ═════════════════════════════════════════════════════════════
    # 6. Conclusion
    # ═════════════════════════════════════════════════════════════
    heading_a("6 Conclusion")
    for text in _section_paras("6"):
        body(text)

    # Acknowledgements omitted for anonymous review.
    # To be added upon acceptance per ACADIA AI disclosure policy.

    # ═════════════════════════════════════════════════════════════
    # References
    # ═════════════════════════════════════════════════════════════
    heading_a("References")
    for ref in REFERENCES:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(12); run.font.name = "Times New Roman"
        pf = p.paragraph_format
        pf.space_after = Pt(3)
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)

    # ── Clear inherited headers/footers ──
    for section in doc.sections:
        for hf in [section.header, section.footer,
                   section.first_page_header, section.first_page_footer]:
            for p in hf.paragraphs:
                if p.text.strip():
                    p.clear()

    # ── Save ──
    doc.save(str(OUTPUT))
    print(f"Paper draft written to {OUTPUT}")

    total = sum(len(t.split()) for _, _, _, pp in SECTIONS for t in pp)
    print(f"Body text: ~{total} words (limit: 4,000)")


def _section_paras(num):
    """Return paragraph list for section with given number."""
    for _, snum, _, paras in SECTIONS:
        if snum == num:
            return paras
    return []


if __name__ == "__main__":
    build()
